import docx
from tkinter import *
from tkinter import scrolledtext, filedialog, messagebox
import io
from pdfminer.converter import TextConverter
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfpage import PDFPage
from datetime import datetime


def choose_clicked():
    global filetype
    filespath_text = StringVar()
    filespath_text.set("")
    file = filedialog.askopenfilename(filetypes=(("Word", "*.docx"), ("Pdf", "*.pdf")))
    filespath_text.set(file)
    filepath = filespath_text.get()
    filetype = filepath[-4:]
    take_file(file)


def start_clicked():
    acad_subj = acad_txt.get()
    teacher_subj = teacher_txt.get()
    if ("".__eq__(acad_subj)) or ("".__eq__(teacher_subj)):
        messagebox.showinfo('Ошибка', 'Не введено Ф.И.О преподавателя \n и (или) название дисциплины.')
    else:
        if len(text) < 1:
            messagebox.showinfo('Ошибка', 'Не выбран файл.')
        else:
            analysis(acad_subj, teacher_subj)


def take_file(file):
    global text
    if filetype == "docx":
        doc = docx.Document(file)
        global image_width
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        for shape in doc.inline_shapes:
            image_width = shape.width
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text.append(cell.text)
    else:
        resource_manager = PDFResourceManager()
        fake_file_handle = io.StringIO()
        converter = TextConverter(resource_manager, fake_file_handle)
        page_interpreter = PDFPageInterpreter(resource_manager, converter)

        with open(file, 'rb') as fh:

            for page in PDFPage.get_pages(fh,
                                          caching=True,
                                          check_extractable=True):
                page_interpreter.process_page(page)

            full_text = fake_file_handle.getvalue()
        text = full_text.split()


def analysis(acad, teacher):
    global image_width
    if filetype == "docx":
        if ((text[0] != "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ") or
                (text[1] != "федеральное государственное автономное образовательное учреждение высшего образования") or
                (text[2] != " «САНКТ-ПЕТЕРБУРГСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ \nАЭРОКОСМИЧЕСКОГО ПРИБОРОСТРОЕНИЯ»") or
                (text[3] != "ИНСТИТУТ НЕПРЕРЫВНОГО И ДИСТАНЦИОННОГО ОБРАЗОВАНИЯ")):
            report.insert(-1, "Ошибка в названии университета.\n")
        if text[4] != "КАФЕДРА КОМПЬЮТЕРНЫХ ТЕХНОЛОГИЙ И ПРОГРАММНОЙ ИНЖЕНЕРИИ ":
            report.insert(-1, "Ошибка в названии кафедры.\n")
        if text[18] != ("Санкт-Петербург " + str(datetime.now().year)):
            report.insert(-1, "Ошибка в указании даты.\n")
        if table_text[4] != teacher:
            report.insert(-1, "Ошибка в имени преподавателя.\n")
        if table_text[12] != ("по дисциплине: " + acad):
            report.insert(-1, "Ошибка в названии дисциплины.\n")
        if table_text[13] != 'СТУДЕНТ(КА)  ГР. №':
            if (table_text[13] == 'СТУДЕНТ  ГР. №') and (table_text[18][-1] == 'a'):
                report.insert(-1, "Ошибка в указании пола студента.\n")
            if (table_text[13] == 'СТУДЕНТКА  ГР. №') and (table_text[18][-1] != 'a'):
                report.insert(-1, "Ошибка в указании пола студента.\n")
        else:
            report.insert(-1, "Не указан пол студента.")
        if image_width == 0:
            report.insert(-1, "Не добавлен скриншот-пример выполнения программы.\n")
    else:
        if text[0:24] != (['МИНИСТЕРСТВО', 'НАУКИ', 'И', 'ВЫСШЕГО',
             'ОБРАЗОВАНИЯ', 'РОССИЙСКОЙ', 'ФЕДЕРАЦИИ',
             'федеральное', 'государственное', 'автономное',
             'образовательное', 'учреждение', 'высшего', 'образования',
             '«САНКТ-ПЕТЕРБУРГСКИЙ', 'ГОСУДАРСТВЕННЫЙ', 'УНИВЕРСИТЕТ',
             'АЭРОКОСМИЧЕСКОГО', 'ПРИБОРОСТРОЕНИЯ»', 'ИНСТИТУТ',
                  'НЕПРЕРЫВНОГО', 'И', 'ДИСТАНЦИОННОГО', 'ОБРАЗОВАНИЯ']):
            report.insert(-1, "Ошибка в названии университета.\n")
        if text[24:30] != (['КАФЕДРА','КОМПЬЮТЕРНЫХ', 'ТЕХНОЛОГИЙ', 'И', 'ПРОГРАММНОЙ', 'ИНЖЕНЕРИИ']):
            report.insert(-1, "Ошибка в названии кафедры.\n")
        if text[79] != str(datetime.now().year):
            report.insert(-1, "Ошибка в указании даты.\n")
        if (text[34] + " " + text[35]) != teacher:
            report.insert(-1, "Ошибка в имени преподавателя.\n")
        if (text[57] + " " + text[58]) != acad:
            report.insert(-1, "Ошибка в названии дисциплины.\n")
        if text[61] != 'СТУДЕНТ(КА)':
            if (text[61] == 'СТУДЕНТ  ГР. №') and (text[67][-1] == 'a'):
                report.insert(-1, "Ошибка в указании пола студента.\n")
            if (text[61] == 'СТУДЕНТКА  ГР. №') and (text[67][-1] != 'a'):
                report.insert(-1, "Ошибка в указании пола студента.\n")
        else:
            report.insert(-1, "Не указан пол студента.")
        for image in text:
            if image == 'Рис.':
                image_width = 1
        if image_width == 0:
            report.insert(-1, "Не добавлен скриншот-пример выполнения программы.\n")
    if len(report) > 0:
        rep_txt = scrolledtext.ScrolledText(window,
                                            width=50, height=10, font=("Times Normal", 12))
        for i in range(len(report)):
            rep_txt.insert(INSERT, report[i])
        rep_txt.place(relx=.5, rely=.7, anchor="center")
    else:
        rep_txt = scrolledtext.ScrolledText(window,
                                            width=60, height=10, font=("Times Normal", 12))
        rep_txt.insert(INSERT, "Ошибки в отчете отсутствуют.")
        rep_txt.place(relx=.5, rely=.7, anchor="center")


text = []
table_text = []
report = []
image_width = 0
filetype = ''
window = Tk()
window.title("Анализ отчета студента")
window.geometry('570x400')
acad_lbl = Label(window, text="Название дисциплины", font=("Times Normal", 12))
acad_lbl.grid(column=0, row=0, ipadx=10, pady=5)
acad_txt = Entry(window, width=30, font=("Times Normal", 12))
acad_txt.grid(column=1, row=0, ipadx=10, pady=5)
teacher_lbl = Label(window, text="И.О. Фамилия преподавателя", font=("Times Normal", 12))
teacher_lbl.grid(column=0, row=1, ipadx=10, pady=5)
teacher_txt = Entry(window, width=30, font=("Times Normal", 12))
teacher_txt.grid(column=1, row=1, ipadx=10, pady=5)
choose_btn = Button(window, text="Выбрать файл", font=("Times Normal", 12), command=choose_clicked)
choose_btn.place(relx=.5, rely=.22, anchor="center")
start_btn = Button(window, text="Начать анализ отчёта", font=("Times Normal", 12), command=start_clicked)
start_btn.place(relx=.5, rely=.33, anchor="center")
window.mainloop()
